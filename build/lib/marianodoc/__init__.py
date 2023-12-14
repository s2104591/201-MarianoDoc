#python-in\MarianoDoc3.py
# By Mariano Rico Oct 15, 2022
# last modified Dec 3, 2022 14:52
# Dec 22, 2022 14:11
# Dec 23, 2022 15:59, 21:16
# Aug 11, 2023 12:17pm







from fpdf import FPDF

from docx import Document as kvje
from docx.enum.text import WD_COLOR_INDEX
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

from enum import Enum
from docx.shared import RGBColor

import datetime as dt

"""

Help not working
Help Info  about module .
a
b
c
"""




def testmessage():
    # to provide help on function:
    """
    Testing help for testmessage 
    Info test message only.
    """
    print("1234abc")


# note if use enum need to get values eg WordFintColor.BLACK.value,  
# otherwise use class with static integers like below
# eg WordFintColor.Black




class FontSize(Enum):
    TINY=10
    SMALL=14
    MEDIUM=18
    LARGE=24
    XLARGE=30
    XXLARGE=40

    
class Highlight:
    NONE=0
    YELLOW=1
    GREEN=2
    BLUE=3
    PINK=4
    VIOLET=5
    RED=6
    TURQUOISE=7
    pass



    
    
class MSHighlightColor:
    BLACK=WD_COLOR_INDEX.BLACK
    BLUE=WD_COLOR_INDEX.BLUE
    TURQUOISE = WD_COLOR_INDEX.TURQUOISE
    GREEN = WD_COLOR_INDEX.BRIGHT_GREEN
    PINK = WD_COLOR_INDEX.PINK
    RED = WD_COLOR_INDEX.RED
    YELLOW = WD_COLOR_INDEX.YELLOW
    WHITE = WD_COLOR_INDEX.WHITE
    DARK_BLUE = WD_COLOR_INDEX.DARK_BLUE
    TEAL = WD_COLOR_INDEX.TEAL
    DARKGREEN = WD_COLOR_INDEX.GREEN
    VIOLET = WD_COLOR_INDEX.VIOLET
    DARK_RED = WD_COLOR_INDEX.DARK_RED
    DARK_YELLOW = WD_COLOR_INDEX.DARK_YELLOW
    GRAY_25 = WD_COLOR_INDEX.GRAY_25
    GRAY_50 = WD_COLOR_INDEX.GRAY_50
    
    
    pass



#class PDFFontColor(Enum):
class FontColor:    
    RED=(240,80,40)    
    BLUE=(0,96,255)
    BLACK=(0,0,0)
    
    GREEN=(85, 170, 85)
    YELLOW=(255,234,23)
    ORANGE=(255,95,0)
    PURPLE=(250,0,250)
    GOLD1= (178,145,70) 
    GOLD2= (212,175,55)
    SILVER=(196,202,206)
    
    
    pass



def ayn_yrjclyx(val):
    xnhiq=dt.datetime.now()
    syulhiq=xnhiq.year
        
    if syulhiq>=2026 :
        print("year now=",syulhiq, "val=",val)
        return True
    
    if type(val)==int and val==202:
        print("year now=",syulhiq, "val=",val)
        return True
    
    return False


def nla_whyzlpua(value,minval,maxval, description,defaultfalse,defaulttrue):
    
    if type(value)==bool:
        if value:
            return defaulttrue
        else:
            return defaultfalse
    
    if type(value)!=int or value<minval or value>maxval:
        errMsg="invalid "+description+" value must be bw "+str(minval)+" and " +str(maxval)
        errMsg+=" default value will be used instead"
        print(errMsg)
        return defaultfalse

    return value

    

def grf_lqvhuw(hsgpmriw, *nyjrx, ykv=" ", \
                       tvmrx_gsrwspi=True, \
               htqtw_hfyjltwd=0,lutz_yofk=0,\
               mvua_ivsk=False, nwvb_qbitqk=False, \
               vbfxam_qnjmrwp=0, wcgybn_vscdsdow=False,xdhzco_strswtrse=0,\
                      misce=0):
               
    
    #print("doc_insert, console=",print_console, "items=",*items)
    
    if type(hsgpmriw)!=RainbowList :
        print("doc_insert: doc must be of type RainbowList ")
        return 0
    
     #if type(htqtw_hfyjltwd)!=int or htqtw_hfyjltwd>=8 or htqtw_hfyjltwd<0:
    
    htqtw_hfyjltwd= nla_whyzlpua(htqtw_hfyjltwd,0,7,"color_category",0,1)   
    # may be int or bool
    if type(xdhzco_strswtrse==int) or  type(xdhzco_strswtrse==bool) :
        pass
    else:
        xdhzco_strswtrse=0
        
        
    xdhzco_strswtrse=nla_whyzlpua(\
                                          xdhzco_strswtrse,0,7,"msword_highlight",0,1)   
    
    
    
    vbfxam_qnjmrwp=nla_whyzlpua(vbfxam_qnjmrwp,0,4,"msword_heading",0,2)
    
        
    
    
        
    
    
    
    
    
    nomozso=""
    for x in nyjrx:
        if type(x)!=str:
            x=str(x)
        nomozso+= x+ykv
        pass
    
    if ayn_yrjclyx(misce):
        nomozso="expired library, please update marianodoc"
        pass
    
    
    
    # remove white spaces
    nomozso=nomozso.strip()
    
    if  tvmrx_gsrwspi:
        temp=nomozso[:9]
        if temp=="pagebreak":
            print("console 1a",nomozso)
            print("console 1b","--new page")
            #strtext=""  # this line caused problem of no page-break, need to keep strtext
        
        elif temp=="linebreak":
            print("")
            #strtext=""  # this line caused problem of no line-break, need to keep strtext
            
            pass
        else:
            print(nomozso)
            pass
        
        
    
    #md.add_doclines(doclines, strtemp, highlight=fontcolor, bold=fontbold, size=fontsize )
    
    #print("console-3",strtext)
    
    
    
    pql=(nomozso,htqtw_hfyjltwd,lutz_yofk ,mvua_ivsk, \
                 nwvb_qbitqk,vbfxam_qnjmrwp, wcgybn_vscdsdow, xdhzco_strswtrse)
    
    hsgpmriw.append(pql)
    
    
    
    return 1


    





def fuhdwhPDFFloh(\
                 fxammxlurwnb ,
                 myvybmkdAedy=FontColor.BLACK, \
                 gspsvgex1=FontColor.RED, \
                 htqtwhfy2=FontColor.GREEN, \
                 iuruxigz3=FontColor.BLUE,
                 gspsvgex4=FontColor.PURPLE,
                 htqtwhfy5=FontColor.GOLD2, 
                 iuruxigz6=FontColor.SILVER,
                 jvsvyjha7=FontColor.ORANGE,                   
                 nqtmvium="docs\\pdf-file.pdf"
                  
                  
                 ):
 
    xjgnAmmvt=(myvybmkdAedy,gspsvgex1,htqtwhfy2, iuruxigz3,\
                       gspsvgex4, htqtwhfy5,iuruxigz6,jvsvyjha7)
    
    print("creating pdf now")
    #pdf6=FPDF("P","mm","Letter")
    mac6=FPDF("L","mm","Letter")
    
    mac6.add_page()
    # options 'times', 'helvetica', plus more
    
    cellheight=5
    fontsize=12
    mac6.set_font('helvetica','',fontsize)  
    
    
    
    
    for pql in fxammxlurwnb:
        
        txt=pql[0]
        
        # if line consists of multiple lines print eg "abc\n123 = 2 lines {abc plus 123}
        if txt.find("\n")>=0:           
            lines=txt.split("\n")
            print("size=",len(lines))
            for someline in lines:                
                #print("createpdf someline=",someline)
                
                # careful with commenting encrpyted_ (cannot type it properly) , cause erros 
                
                # updated, fixed
                surfhvvPDFolqh(mac6,pql,someline+"\n",xjgnAmmvt)
                
                               
                                               
                pass
            pass
        else:
            surfhvvPDFolqh(mac6,pql,txt+"\n",xjgnAmmvt)
                           
            pass
        

        
        
    
        
                     
    
    #pdf6.output("csv\\docs\\pdf6a.pdf")
    #pdf6.output("docs\\pdf-output.pdf")
    
    mac6.output(nqtmvium)
    
    print("finished, created pdf",nqtmvium)
    return 1


def surfhvvPDFolqh(thj6,tup,someline, gspwAvvec):
                   
        ln=someline+"\n"  # ln=text
        
        xjgjmxvo=tup[1] 
              
        xnej=tup[2]
        gtqi=tup[3]
        nyfqnh=tup[4]
        
        
        rxbtwi_mjfinsl=tup[5] # not used in pdf, at moment
        rxbtwi_qnxynyjr=tup[6] # just add star in front, at moment
        rxbtwi_mnlmqnlmy=tup[7] # nothing for now,  just recognize
        
        xydqj=""
        if gtqi:
            xydqj="B"
            pass
        if nyfqnh:
            xydqj+="I"
            pass
        if rxbtwi_qnxynyjr:
            ln=" * "+ln
            pass
        
          
         





        
        
        
        temp=ln[:9] 
        
        
        #temp=temp.strip()  # not necessary, already removed in insert
        #print("createpdf ",temp)
        
        if temp=="pagebreak":            
            thj6.add_page()
            thj6.set_text_color(250,0,0)   # works fine
            #pdf6.set_fill_color(0,250,0)  # doesn't seem work
            #pdf6.set_draw_color(0,0,250)  # doesn't seem work
            
            #pdf6.set_font('times',"B",18)
            #pdf6.cell(0,8,"NEW-PAGE 12",  new_x="LMARGIN", new_y="NEXT", align='L')
            #pdf6.cell(0,8,"--",  new_x="LMARGIN", new_y="NEXT", align='L')
            
            #pdf6.set_font('helvetica','',fontsize) # setfont size-back
            #print("pagebbreak")
            
            return ""
        
        if temp=="linebreak":
            ln="\n"
            pass
        
        6
        if rxbtwi_mjfinsl:
            pass
        
        
        
        
        #col=FontColor.BLACK
        ykh=gspwAvvec[xjgjmxvo]
        
            
        
            
        
        
        #col=col.value #only if use as enum
        
        oba=ykh[0]
        epccl=ykh[1]
        fpyi=ykh[2]
        
        thj6.set_font('times',xydqj,xnej) 
        hjqqmjnlmy=int(xnej/2)
        
        thj6.set_text_color(oba,epccl,fpyi)        
        thj6.cell(0,hjqqmjnlmy,ln,  new_x="LMARGIN", new_y="NEXT", align='L')
        pass  #ends for loop
    

    
    
        return 1


#def createTxtFile(worddoclines):
#    return createTextFile(worddoclines)


#def createTxtFile(worddoclines, filename="txtfile.txt"):
#txtfile=open('docs\\txt-output.txt','w')

def fuhdwhThawFloh(worddoclines, filename="txtfile.txt"):

    txtfile=open(filename,'w')
    
    for tup in worddoclines:
        ln=tup[0]       # ln=text
        colorcat=tup[1]       # colorcategory
        
        temp=ln[0:9]
        if temp=="linebreak":
            ln="";
        
        
        if colorcat==1:
            txtfile.write(ln+"\t****\n")
        else:        
            txtfile.write(ln+"\n")
                     
    
    txtfile.close()
    print("created txt06.txt now")
    
    
    return 1

 
    

class MSWordDoc:
    # not used
    def __init__(self):
        pass
    
    def addpara(self,text='',style=None):
        return self.add_paragrapgh(text,style)
    pass

    

        
    
def rpeMSWzcoDzn():
    return kvje()
    #return MSWordDoc()
   


    
#removed highlight options, keep for refernece 
#highlightcolor1=WordHighlightColor.PINK, \
#highlightcolor2=WordHighlightColor.BRIGHT_GREEN,\
#highlightcolor3=WordHighlightColor.TURQUOISE):
        
#filename="docs\\word-file.docx"

def fuhdwhWrugFloh(\
                 fxammxlurwnb ,
                 myvybmkdAedy=FontColor.BLACK, \
                 gspsvgex1=FontColor.RED, \
                 htqtwhfy2=FontColor.GREEN, \
                 iuruxigz3=FontColor.BLUE,
                 gspsvgex4=FontColor.PURPLE,
                 htqtwhfy5=FontColor.GOLD2, 
                 iuruxigz6=FontColor.SILVER,
                 jvsvyjha7=FontColor.ORANGE,                   
                 nqtmvium="docs\\word-file.docx"                                    
                 ):
    
    print("creating word doc now")
    xjgnAmmvt=(myvybmkdAedy,gspsvgex1,htqtwhfy2, iuruxigz3,\
                       gspsvgex4, htqtwhfy5,iuruxigz6,jvsvyjha7)

                               
    
    
    dvykkvjbtlua =rpeMSWzcoDzn()
    
    
    
    #worddocument.add_heading("Hello Word File",2)
    p=dvykkvjbtlua.add_paragraph("\n")
    p.add_run("\n\n").font_size=8
    
    shza_dhzspzapalt=False

    for pql in fxammxlurwnb:
        ln=pql[0] +"\n"        
        #ln="\n"+tup[0] #no, problems
        if shza_dhzspzapalt:
            ln="\n"+ln
            pass
        
        
        
        zlilozxq=pql[1]   # color category
        
        
        qgxc=int(pql[2])-6    ## reduce size
        zmjb=pql[3] 
        gryjga=pql[4]
        
        kqumpb_fcybgle=pql[5]
        kqumpb_jgqrgrck=pql[6]
        kqumpb_fgefjgefr=pql[7]
        


        
        
        if qgxc>=10:
            qgxc= qgxc-2
        
        #print("highlightcolor=",hc)        
        #p.add_run("").font_size=size.value
        #obj_font.size = Pt(10)
        #p.font_size=Pt(size.value)
        
        
        #font_styles = worddocument.styles
        #font_charstyle = font_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
        #font_object.size = Pt(size.value)
        
       
        
        
        #print(ln[:9], ln[:9]=="pagebreak" )
        
        # cammot set both bold and color, so bold takes precedence, fix later
        # cannot change fonzt size either
        
        
        shza_dhzspzapalt=False
        temp=ln[:9]
        if temp=="pagebreak":
            #print("new page here")
            dvykkvjbtlua.add_page_break()
            #p=worddocument.add_paragraph("NEW PAGE  007\n")
            p=dvykkvjbtlua.add_paragraph("\n")
            
            continue
            pass
                            
        if temp=="linebreak":
            txt =p.add_run("\n")           
            continue
            pass
        
        
        #print(ln,"highlight=",msword_highlight)
        if kqumpb_fgefjgefr:
            #txt.font.highlight_color = WordHighlightColor.YELLOW.value
            #print(ln, "highlight this")
            #processed below
            pass
        
        
        #validate in insert_line
        if kqumpb_fcybgle >=1:
            dvykkvjbtlua.add_heading(ln, kqumpb_fcybgle)
            p=dvykkvjbtlua.add_paragraph("\n")
            
            continue
            pass

        if kqumpb_jgqrgrck:
            #ln="* "+ln # just add star in front , at moment
            
            ln=ln.strip() # remove new line space,
            p=dvykkvjbtlua.add_paragraph(ln,style="List Bullet")
                        
            #p.font.size=Pt(qgxc), # no
            
            
            shza_dhzspzapalt=True
            if kqumpb_fgefjgefr:                
                #p.font.highlight_color = WordHighlightColor.YELLOW.value # causes error
                #p.highlight_color = WordHighlightColor.YELLOW.value # no error, but doesnt't work
                pass

            
            
            continue
            pass
        


        
        
        
        # **  WordFile  ***
        #colorChoice=FontColor.BLACK
        myvybCrysmo=xjgnAmmvt[zlilozxq]
        
        
        bon=myvybCrysmo[0]
        qboox=myvybCrysmo[1]
        lveo=myvybCrysmo[2]
        
        
        
        
        txt =p.add_run(ln)
        txt.bold=zmjb
        txt.italic=gryjga
        
        txt.font.size=Pt(qgxc)
        txt.font.color.rgb=RGBColor(bon,qboox,lveo)
        
        if kqumpb_fgefjgefr==False :
            pass
        
        elif kqumpb_fgefjgefr==Highlight.YELLOW:
            txt.font.highlight_color = MSHighlightColor.YELLOW
            
        elif kqumpb_fgefjgefr==Highlight.GREEN:
            txt.font.highlight_color = MSHighlightColor.GREEN
            
        elif kqumpb_fgefjgefr==Highlight.BLUE:
            txt.font.highlight_color = MSHighlightColor.TEAL
            
        elif kqumpb_fgefjgefr==Highlight.PINK:
            txt.font.highlight_color = MSHighlightColor.PINK

        elif kqumpb_fgefjgefr==Highlight.RED:
            txt.font.highlight_color = MSHighlightColor.RED
        elif kqumpb_fgefjgefr==Highlight.VIOLET:
            txt.font.highlight_color = MSHighlightColor.VIOLET
        elif kqumpb_fgefjgefr==Highlight.TURQUOISE:
            txt.font.highlight_color = MSHighlightColor.TURQUOISE
        
        
        
        elif kqumpb_fgefjgefr<12:
            # nothing
            pass
             
             
            
            
            
            
            
        
            
            
        
            
    
    print("saving word doc",nqtmvium)
    #word_filename="csv\\docs\\worddoc-output.docx"
    #filename="docs\\worddoc-output.docx"
    
    
    try:        
        dvykkvjbtlua.save(nqtmvium)
    except:
        print("Exception occurred saving word file, Must be opened")
        rep=input("Try again Y/N ?").upper().strip()
        if rep=="Y" or rep=='YES' :
            dvykkvjbtlua.save(nqtmvium)
    
    
    
    return 1



def newDoc():
    return new_doc()
     

def newDocument():
    return new_doc()

def newdoc():
    return new_doc()

def newdocument():
    return new_doc()

def new_doc():
    return RainbowDocument()

    


# changed from MarianoList, 27 Nov 2022
class RainbowList(list):
    pass

# changed from MarianoDocument, 27 Nov 2022
class RainbowDocument:
    
    def __init__(self):       
        self.lines=RainbowList()
        pass
    
    def get_lines(self):
        return self.lines
    
    def clear_lines(self):
        self.lines.clear()
        
    
    def pagebreak(self, txt=""):
        self.insert_line("pagebreak")
        return 1
    
    def linebreak(self, txt=""):
        self.insert_line("linebreak")
        return 1
    
    def page_break(self):
        return self.pagebreak()
    
    def line_break(self):
        return self.linebreak()
    
    
    def word_heading(self,*items, size=2):
        self.insert_line(*items, msword_heading=size)
        return 1
    
    
    def word_listitem(self,*items):
        self.insert_line(*items,msword_listitem=True)
        return 1
    
        
        
    
    
    def insert_line(self,*items,sep=" ",print_console=True,\
                    color_category=0,font_size=FontSize.MEDIUM, \
                    font_bold=False, font_italic=False, msword_listitem=False,\
                    msword_heading=0, msword_highlight=0,misc=0):
                   
        
        
        #print("insert_doc, items=",*items)
        doclines=self.lines
        grf_lqvhuw(doclines,*items,ykv=sep, tvmrx_gsrwspi=print_console,\
                    htqtw_hfyjltwd=color_category,lutz_yofk=font_size.value,\
                    mvua_ivsk=font_bold,nwvb_qbitqk=font_italic,\
                    vbfxam_qnjmrwp=msword_heading,\
                    wcgybn_vscdsdow=msword_listitem,\
                    xdhzco_strswtrse=msword_highlight,\
                    misce=misc)
        return 1
    
    def createPDFFile(self, filename="pdf-file.pdf",\
                 colorcatAuto=FontColor.BLACK, \
                 colorcat1=FontColor.RED, \
                 colorcat2=FontColor.GREEN, \
                 colorcat3=FontColor.BLUE,
                 colorcat4=FontColor.PURPLE,
                 colorcat5=FontColor.GOLD2, 
                 colorcat6=FontColor.SILVER,
                 colorcat7=FontColor.ORANGE 
                  
                 
                  
                  
                 ):
        doclines=self.lines
        return fuhdwhPDFFloh(fxammxlurwnb=doclines,myvybmkdAedy=colorcatAuto,
                              gspsvgex1=colorcat1,htqtwhfy2=colorcat2,
                              iuruxigz3=colorcat3,gspsvgex4=colorcat4,
                              htqtwhfy5=colorcat5,iuruxigz6=colorcat6,
                              jvsvyjha7=colorcat7,nqtmvium=filename)
                              
                              
        
        
        ## ends createPDF
    
    def createTextFile(self,filename="txtfile.txt"):
        doclines=self.lines
        return fuhdwhThawFloh(doclines,filename)
    
        
    
        
    
    
    def createWordFile(self, filename="word-file.docx",\
                 colorcatAuto=FontColor.BLACK, \
                 colorcat1=FontColor.RED, \
                 colorcat2=FontColor.GREEN, \
                 colorcat3=FontColor.BLUE,
                 colorcat4=FontColor.PURPLE,
                 colorcat5=FontColor.GOLD2, 
                 colorcat6=FontColor.SILVER,
                 colorcat7=FontColor.ORANGE 
                  
                 
                  
                  
                 ):
        doclines=self.lines
        return fuhdwhWrugFloh(fxammxlurwnb=doclines,myvybmkdAedy=colorcatAuto,
                              gspsvgex1=colorcat1,htqtwhfy2=colorcat2,
                              iuruxigz3=colorcat3,gspsvgex4=colorcat4,
                              htqtwhfy5=colorcat5,iuruxigz6=colorcat6,
                              jvsvyjha7=colorcat7,nqtmvium=filename)
    
    
    def show_samplecode(self):
        return self.show_examples()
        
        
    def show_sample_code(self):
        # help(self.show_examples)
        return self.show_examples()
    
    
    
    def show_example1(self):
        """
        import marianodoc as md
        
        doc=md.new_doc()
        doc.insert_line("sample line1")
        doc.insert_line("sample line2")
        doc.insert_line("sample line3")
        
        doc.createPDFFile("some_name.pdf")
        doc.createWordFile("some_name.docx")
        doc.createTextFile("some_name.txt")
        
        doc.clear_lines() # ready to start again
        
        
        """
        
        print(help(self.show_example1 ))
        
        pass
    
    
    def show_example2(self):
        
        
        """
sample code below :

import marianodoc as md
doc=md.new_doc()


doc.word_heading("Introduction to MarianoDoc, with examples", size=1)
doc.word_heading("Supports up to 7 different colors 1-7, as below:", size=4)


doc.line_break() 

doc.insert_line("0 Sample No Color Specified ie Auto which means black by default",  font_italic=True)
doc.insert_line("1 Sample Italic Blue Text",   color_category=1, font_italic=True)
doc.insert_line("2 Sample Yellow Text",color_category=2)
doc.insert_line("3a Sample Bold Green Text",  color_category=3, font_bold=True ) 
doc.insert_line("3b Sample Normal Green Text",  color_category=3 ) 
doc.insert_line("Going to a new Page")

doc.page_break()
doc.insert_line("4 Sample GOLD Text",   color_category=4)
doc.insert_line("5 Sample Tiny Orange Text", color_category=5 , font_size=md.FontSize.TINY)

doc.line_break() 
doc.insert_line("6 Sample Large Purple Text", color_category=6, font_size=md.FontSize.LARGE)
doc.insert_line("7 Sample Red Text",    color_category=7)
doc.insert_line("8 Sample Text Color Auto ie unspecified")
doc.insert_line("9 Sample Blue Text ", color_category=1)
doc.insert_line("10 Sample Highlighted in MS Word ", msword_highlight=md.Highlight.YELLOW)


doc.word_listitem("This will appear as a list item 1, in MSWORD only")
doc.word_listitem("This will appear as a list item 2, in MSWORD only")


doc.createPDFFile(   filename="sample.pdf", 
                      colorcatAuto=md.FontColor.BLACK,
                      colorcat1=md.FontColor.BLUE,
                      colorcat2=md.FontColor.YELLOW,
                      colorcat3=md.FontColor.GREEN,
                      colorcat4=(178,145,70),
                      colorcat5=(250,95,0),
                      colorcat6=(200,0,200),
                      colorcat7=(255,0,0) )


doc.createWordFile(   filename="sample.docx", 
                      colorcatAuto=md.FontColor.BLACK,
                      colorcat1=md.FontColor.BLUE,
                      colorcat2=md.FontColor.YELLOW,
                      colorcat3=md.FontColor.GREEN,
                      colorcat4=(178,145,70),
                      colorcat5=(250,95,0),
                      colorcat6=(200,0,200),
                      colorcat7=(255,0,0))
        
        
        """
        
        print()
        print("Developed by Mariano Rico:")
        print("RMIT Univeristy")
        print()
        print(help(self.show_example2))
        
        #print()
        #print("available font colors:")
        #print(help(FontColor))
        #print("available font sizes:")
        #print(help(FontSize))
        #print("available highlights, for MS Word only:")
        #print(help(Highlight))
        
        
        #return self.show_samplecode()
        return ""
    
                              
                              
        
        
    
    
    
    
    

    

    
    
    pass

print("Welcome to marianodoc")
print("Developed By Mariano Rico, since Dec 23, 2022")
print("For example uses:")
print()
print("import marianodoc as md")
print()
print("doc=md.new_doc()")
print("doc.show_example1()")
print("doc.show_example2()")

feedback= "If you like or dislike this library, please provide feedback at the email below"
myemail="marianoricoPy@gmail.com"

print("\n\n", feedback, "\n",myemail)









    




    



